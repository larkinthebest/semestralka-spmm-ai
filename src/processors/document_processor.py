import os
from pathlib import Path
from typing import Optional
import pypdf # Changed from PyPDF2
from docx import Document as DocxDocument
import re # Import regex module

class DocumentProcessor:
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx']
    
    def extract_text(self, file_path: str) -> str:
        """Extract text from PDF or DOCX files"""
        file_extension = Path(file_path).suffix.lower()
        
        if file_extension == '.pdf':
            return self._extract_from_pdf(file_path)
        elif file_extension == '.docx':
            return self._extract_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file) # Changed from PyPDF2.PdfReader
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n"
            
            return text.strip()
        except Exception as e:
            print(f"Error extracting text from PDF: {e}")
            return "Error: Could not extract text from PDF file."
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = DocxDocument(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text.strip()
        except Exception as e:
            print(f"Error extracting text from DOCX: {e}")
            return "Error: Could not extract text from DOCX file."
    
    def get_document_info(self, file_path: str) -> dict:
        """Get basic information about the document"""
        file_stats = os.stat(file_path)
        file_extension = Path(file_path).suffix.lower()
        
        info = {
            "filename": Path(file_path).name,
            "size": file_stats.st_size,
            "format": file_extension,
            "supported": file_extension in self.supported_formats
        }
        
        if file_extension == '.pdf':
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = pypdf.PdfReader(file) # Changed from PyPDF2.PdfReader
                    info["pages"] = len(pdf_reader.pages)
            except:
                info["pages"] = "Unknown"
        
        return info
    
    def chunk_text(self, text: str, chunk_size: int = 1000, overlap: int = 100) -> list:
        """Split text into chunks for processing"""
        if len(text) <= chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            
            # Try to break at sentence boundary
            if end < len(text):
                # Look for sentence endings
                for i in range(end, max(start + chunk_size - 200, start), -1):
                    if text[i] in '.!?':
                        end = i + 1
                        break
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            start = end - overlap
        
        return chunks
    
    def clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        # Remove excessive whitespace
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            if line:  # Skip empty lines
                cleaned_lines.append(line)
        
        # Join lines and normalize spacing
        cleaned_text = ' '.join(cleaned_lines)
        
        # Remove multiple spaces
        import re
        cleaned_text = re.sub(r'\s+', ' ', cleaned_text)
        
        return cleaned_text.strip()

    def detect_and_convert_latex(self, text: str) -> str:
        """
        Detects LaTeX formulas in the given text and converts them to a standardized format.
        This method focuses on identifying common LaTeX patterns and ensuring they are
        properly delimited for frontend rendering.
        """
        # Pattern for inline LaTeX: $...$
        # Pattern for display LaTeX: $$...$$ or \[...\] or \begin{...}...\end{...}
        
        # Convert inline $...$ to \(...\)
        text = re.sub(r'\$(.*?)\$', r'\\(\1\\)', text)
        
        # Convert display $$...$$ to \[...\]
        text = re.sub(r'\$\$(.*?)\$\$', r'\\[\1\\]', text)

        # Handle common OCR errors for mathematical notation
        # 1. Replace 'a' with 'x' in derivative contexts (heuristic)
        text = re.sub(r'(\d+)?a([\'"]+)\s*\(t\)', r'\1x\2(t)', text)
        text = re.sub(r'a([\'"]+)\s*\(t\)', r'x\1(t)', text) # For x'(t), x''(t)
        
        # 2. Replace '**' with '^{' and '}' for superscripts
        # This is a more complex heuristic. We'll look for patterns like 'e**3t' and convert to 'e^{3t}'
        text = re.sub(r'([a-zA-Z0-9])\*\*([a-zA-Z0-9]+)', r'\1^{\2}', text)
        
        # 3. Ensure common environments are properly delimited if not already
        # This is a more complex task and might require a dedicated LaTeX parser for full robustness.
        # For now, we'll focus on common delimiters.
        
        # Example: \begin{equation} ... \end{equation}
        # We assume these are already correctly formatted for MathJax/KaTeX.
        
        return text
